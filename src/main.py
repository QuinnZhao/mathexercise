from PyQt5.QtWidgets import  QApplication, QMainWindow, QMessageBox
from math_g1_test import Ui_Form


class MathTestGenWindow (QMainWindow, Ui_Form):
    def __init__ (self):
        super (MathTestGenWindow, self).__init__()
        self.setupUi (self)
        self.btnExport.setEnabled(False)
        self.teItems.setAutoFillBackground(True)
        self.gap = " " * 4
        self.gap = "\t\t"
        self.template = "{1}{0}{2}{0}{3}{0}{4}{0}{5}"
        self.add_template = "{}+{}+{}="
        self.sub_template = "{}-{}-{}="
        self.add_sub_template = "{}+{}-{}="
        self.sub_add_template = "{}-{}+{}="
        item1 = "2+6+1="
        item2 = "9-5-4="
        item3 = "6+3+1="
        item4 = "2+7+1="
        item5 = "9-5-4="
        self.items_list = []
        self.teItems.setText(self.template.format(self.gap, item1, item2, item3, item4, item5))

    def generate(self):
        """        '''
        根据所选题型随机生成10以内的口算题目
        return: None
        """
        if self.cbbType.currentText() == "10以内连加":
            self.gen_add()
        elif self.cbbType.currentText() == "10以内连减":
            self.gen_sub()
        elif self.cbbType.currentText() == "10以内加减混合":
            self.gen_mix()
        self.output_to_te()
        self.btnExport.setEnabled(True)

    def gen_add(self):
        import random
        self.items_list = []
        count = 0
        while count <= int(self.cbbNum.currentText()):
            while True:
                first_num = random.randint(0,9)
                second_num = random.randint(0,9)
                third_num = random.randint(0,9)
                if first_num + second_num + third_num <= 10:
                    self.items_list.append(self.add_template.format(first_num,second_num,third_num))
                    count += 1
                    break
        # print(self.items_list)

    def output_to_te(self):
        text = ""
        for i in range(0, int(self.cbbNum.currentText()), 5):
            text += self.template.format(self.gap,
                self.items_list[i], self.items_list[i+1], self.items_list[i+2],
                self.items_list[i+3], self.items_list[i+4])
            text += "\n"

        self.teItems.setText(text)

    def gen_sub(self):
        import random
        self.items_list = []
        count = 0
        while count <= int(self.cbbNum.currentText()):
            while True:
                first_num = random.randint(0,10)
                second_num = random.randint(0,9)
                third_num = random.randint(0,9)
                if first_num - second_num - third_num >= 0:
                    self.items_list.append(self.sub_template.format(first_num,second_num,third_num))
                    count += 1
                    break
        # print(self.items_list)

    def gen_mix(self):
        import random
        self.items_list = []
        count = 0
        while count <= int(self.cbbNum.currentText()):
            while True:
                first_num = random.randint(0, 10)
                second_num = random.randint(0, 9)
                third_num = random.randint(0, 9)
                if first_num + second_num + third_num <= 10:
                    self.items_list.append(self.add_template.format(first_num, second_num, third_num))
                    count += 1
                    break
                elif first_num - second_num - third_num >= 0:
                    self.items_list.append(self.sub_template.format(first_num, second_num, third_num))
                    count += 1
                    break
                elif (0 <= first_num + second_num - third_num <= 10) and (first_num + second_num <= 10):
                    self.items_list.append(self.add_sub_template.format(first_num, second_num, third_num))
                    count += 1
                    break
                elif (0 <= first_num - second_num + third_num <= 10) and (first_num - second_num >= 0):
                    self.items_list.append(self.sub_add_template.format(first_num, second_num, third_num))
                    count += 1
                    break

    def export(self):
        from PyQt5.QtWidgets import QFileDialog

        filename, _ = QFileDialog.getSaveFileName(self, "文件保存", "./",
                        "所有文件 (*);;文本文件 (*.txt);; Word文件 (*.docx);; CSV 文件 (*.csv)")
        file_type = filename.split(".")[-1]

        if file_type.lower() == "txt":
            self.export_to_text(filename)

        if file_type.lower() == "csv":
            self.export_to_csv(filename)

        if file_type == "docx":
            self.export_to_word(filename)

    def export_to_text(self, filename):
        with open(filename, "w") as f:
            f.write("\n\n\t\t一年级{}口算练习题".format(self.cbbType.currentText()))
            f.write("\n")
            f.write("\t\t姓名：   时间：  做对：（   ）题\n\n")
            f.write(self.teItems.toPlainText())

    def export_to_csv(self, filename):
        with open(filename, "w") as f:
            for i in range(0, int(self.cbbNum.currentText()), 5):
                f.write(self.items_list[i] + ",")
                f.write(self.items_list[i + 1] + ",")
                f.write(self.items_list[i + 2] + ",")
                f.write(self.items_list[i + 3] + ",")
                f.write(self.items_list[i + 4] + "\n")

    def export_to_word(self, filename):
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        doc = Document("./templates/default.docx")
        paragraph = doc.add_paragraph("一年级{}口算练习题\n".format(self.cbbType.currentText()))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.bold = True
        paragraph = doc.add_paragraph("姓名：\t\t时间：\t\t做对：（   ）题\n")
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table = doc.add_table(rows=int(self.cbbNum.currentText())//5, cols=5)
        for i in range(0, int(self.cbbNum.currentText()), 5):
            table.cell(i // 5, 0).text = self.items_list[i]
            table.cell(i // 5, 1).text = self.items_list[i + 1]
            table.cell(i // 5, 2).text = self.items_list[i + 2]
            table.cell(i // 5, 3).text = self.items_list[i + 3]
            table.cell(i // 5, 4).text = self.items_list[i + 4]

        doc.save(filename)

if __name__ == "__main__":
    import sys
    app = QApplication(sys.argv)
    window = MathTestGenWindow()
    window.show()
    sys.exit(app.exec_())
